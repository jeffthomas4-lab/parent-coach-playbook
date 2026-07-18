from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path('strategy/PARENT-COACH-PLAYBOOK-COMPANY-OPERATING-SYSTEM.md')
OUTPUT = Path('strategy/Parent_Coach_Playbook_Company_Operating_System.docx')

NAVY = RGBColor(22, 50, 68)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
RUST = RGBColor(151, 65, 43)
GRAY = RGBColor(91, 101, 112)
LIGHT_GRAY = 'D9E0E6'


def set_run(run, *, font='Calibri', size=11, color=None, bold=None, italic=None):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), font)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), font)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for edge, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        tag = f'w:{edge}'
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement('w:tblHeader')
    repeat.set(qn('w:val'), 'true')
    tr_pr.append(repeat)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fld_char, instr, separate, text, end])


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in [
        ('Heading 1', 16, BLUE, 18, 10),
        ('Heading 2', 13, BLUE, 14, 7),
        ('Heading 3', 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ('List Bullet', 'List Number'):
        style = styles[name]
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if 'Kicker' not in styles:
        kicker = styles.add_style('Kicker', WD_STYLE_TYPE.PARAGRAPH)
    else:
        kicker = styles['Kicker']
    kicker.font.name = 'Calibri'
    kicker.font.size = Pt(10)
    kicker.font.bold = True
    kicker.font.color.rgb = RUST
    kicker.paragraph_format.space_after = Pt(10)

    header = section.header.paragraphs[0]
    header.text = 'PARENT COACH PLAYBOOK  |  COMPANY OPERATING SYSTEM'
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.runs[0], size=8.5, color=GRAY, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run('Internal operating document  |  ')
    set_run(run, size=8.5, color=GRAY)
    add_page_field(footer)


def add_inline(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|“.*?”)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            set_run(run, font='Consolas', size=9.5, color=DARK_BLUE)
        else:
            paragraph.add_run(part)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph(style='Kicker')
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.add_run('COMPANY OPERATING SYSTEM')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run('Parent Coach Playbook')
    set_run(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run('How the company discovers, decides, ships, distributes, maintains, learns, and grows')
    set_run(run, size=14, color=DARK_BLUE, italic=True)

    values = [
        ('Version', '1.0'),
        ('Effective', 'July 18, 2026'),
        ('Owner', 'Jeff Thomas'),
        ('Horizon', 'May 2026 through May 2028'),
    ]
    for label, value in values:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        label_run = p.add_run(f'{label}: ')
        set_run(label_run, size=10.5, color=NAVY, bold=True)
        value_run = p.add_run(value)
        set_run(value_run, size=10.5, color=GRAY)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_contents(doc, headings):
    p = doc.add_paragraph('How to use this document', style='Heading 1')
    p.paragraph_format.space_before = Pt(0)
    add_inline(doc.add_paragraph(), 'Read sections 1–4 for direction, sections 5–13 for the operating system, and sections 14–16 for current execution. Detailed specialist runbooks remain authoritative inside their domains.')
    doc.add_paragraph('Contents', style='Heading 1')
    for heading in headings:
        p = doc.add_paragraph(style='List Number')
        add_inline(p, re.sub(r'^\d+\.\s*', '', heading))
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def parse_blocks(lines):
    blocks = []
    paragraph = []
    for raw in lines:
        line = raw.rstrip()
        if not line:
            if paragraph:
                blocks.append(('paragraph', ' '.join(paragraph)))
                paragraph = []
            continue
        if line.startswith('#'):
            if paragraph:
                blocks.append(('paragraph', ' '.join(paragraph)))
                paragraph = []
            level = len(line) - len(line.lstrip('#'))
            blocks.append((f'heading{level}', line[level:].strip()))
        elif re.match(r'^\d+\.\s+', line):
            if paragraph:
                blocks.append(('paragraph', ' '.join(paragraph)))
                paragraph = []
            blocks.append(('number', re.sub(r'^\d+\.\s+', '', line)))
        elif line.startswith('- '):
            if paragraph:
                blocks.append(('paragraph', ' '.join(paragraph)))
                paragraph = []
            blocks.append(('bullet', line[2:]))
        elif line.startswith('**') and ':**' in line:
            if paragraph:
                blocks.append(('paragraph', ' '.join(paragraph)))
                paragraph = []
            blocks.append(('metadata', line.replace('**', '')))
        else:
            paragraph.append(line)
    if paragraph:
        blocks.append(('paragraph', ' '.join(paragraph)))
    return blocks


def build(source=SOURCE, output=OUTPUT):
    text = source.read_text(encoding='utf-8')
    lines = text.splitlines()
    headings = [line[3:].strip() for line in lines if line.startswith('## ') and re.match(r'\d+\.', line[3:])]
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc, headings)

    blocks = parse_blocks(lines)
    for kind, value in blocks:
        if kind == 'heading1':
            continue
        if kind == 'metadata':
            continue
        if kind.startswith('heading'):
            level = min(int(kind[-1]) - 1, 3)
            p = doc.add_paragraph(style=f'Heading {level}')
            add_inline(p, value)
        elif kind == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, value)
        elif kind == 'number':
            p = doc.add_paragraph(style='List Number')
            add_inline(p, value)
        else:
            p = doc.add_paragraph()
            add_inline(p, value)

    props = doc.core_properties
    props.title = 'Parent Coach Playbook Company Operating System'
    props.subject = 'Company strategy, operating cadence, discovery, maintenance, governance, and growth system'
    props.author = 'Parent Coach Playbook'
    props.keywords = 'company operating system, strategy, discovery, maintenance, scorecard, governance'
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f'Wrote {output}')


if __name__ == '__main__':
    build(Path(sys.argv[1]) if len(sys.argv) > 1 else SOURCE, Path(sys.argv[2]) if len(sys.argv) > 2 else OUTPUT)
